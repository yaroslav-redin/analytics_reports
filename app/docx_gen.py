import io
import time
import random
import re
import json
import hashlib
import threading
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.chart_gen import insert_visualization
from app import config as cfg


# ===================== КЭШ АНАЛИТИКИ В ПАМЯТИ ПРОЦЕССА =====================
# Сбрасывается при рестарте uvicorn; этого достаточно для одной сессии работы
# и для случаев когда пользователь жмёт «экспорт» повторно с теми же данными.

_analysis_cache: dict[str, str] = {}
_analysis_cache_lock = threading.Lock()


def _cache_key(q: dict, sec_name: str, sec_desc: str) -> str:
    payload = json.dumps(
        {
            "q":     q.get("question_name"),
            "rows":  q.get("rows"),
            "sec":   sec_name,
            "desc":  sec_desc,
            "model": cfg.get("openrouter_model"),
            # включаем стиль/правила в ключ — при правке промпта кэш инвалидируется
            "style": hashlib.md5(cfg.get("prompt_style_example").encode()).hexdigest(),
            "rules": hashlib.md5(cfg.get("prompt_writing_rules").encode()).hexdigest(),
        },
        ensure_ascii=False,
        sort_keys=True,
        default=str,
    )
    return hashlib.md5(payload.encode("utf-8")).hexdigest()


# ===================== SINGLETON КЛИЕНТ И PACING =====================

_openrouter_client: OpenAI | None = None
_openrouter_client_key: str = ""
_LAST_REQUEST_TIME: float = 0.0


def get_openrouter_client() -> OpenAI:
    """Один TCP-клиент на всё приложение. Пересоздаётся, если ключ изменился через настройки."""
    global _openrouter_client, _openrouter_client_key
    key = cfg.get("openrouter_api_key")
    if not key:
        raise ValueError("OPENROUTER_API_KEY не задан (проверьте .env или Настройки)")
    if _openrouter_client is None or _openrouter_client_key != key:
        _openrouter_client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=key,
            timeout=cfg.get_float("llm_request_timeout"),
        )
        _openrouter_client_key = key
    return _openrouter_client


def _pace():
    """Адаптивная пауза: ждём только если с прошлого запроса прошло меньше min_interval."""
    global _LAST_REQUEST_TIME
    min_interval = cfg.get_float("llm_sleep_between_calls")
    delta = time.time() - _LAST_REQUEST_TIME
    if delta < min_interval:
        time.sleep(min_interval - delta)


def _mark_request_done():
    global _LAST_REQUEST_TIME
    _LAST_REQUEST_TIME = time.time()


def _backoff_wait(attempt: int):
    """Экспоненциальный backoff с jitter: 2/4/8/16/32с (capped at 60) + random 0.5..2."""
    wait = min(60.0, (2 ** (attempt + 1)) + random.uniform(0.5, 2.0))
    print(f"  -> Rate limit (попытка {attempt + 1}), жду {wait:.1f}с...")
    time.sleep(wait)


# ===================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =====================

def _p(doc, text, bold=False, size=14, space_before=0, space_after=0,
       center=False, first_line_indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(1.25)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return para


def _make_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(14)
    return doc


# ===================== ТИТУЛЬНЫЙ ЛИСТ =====================

def _p_inline_bold(doc, line: str, size: int = 14,
                   space_before: int = 0, space_after: int = 0,
                   center: bool = False, first_line_indent: bool = True):
    """
    Рендерит одну строку как абзац с поддержкой инлайн-разметки **bold**.
    Используется для рендера тела титульного листа: ключи вроде
    «**Основание для проведения исследования** – …» получают жирный фрагмент.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(1.25)

    # **bold** парсится через re.split: чётные сегменты — обычные, нечётные — жирные.
    parts = re.split(r"\*\*(.+?)\*\*", line)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
    return para


def _render_approval_stamp(doc, approval_text: str):
    """
    Рендерит гриф «УТВЕРЖДАЮ» как 1×1 таблицу справа без границ.
    Каждая строка из approval_text — отдельный абзац внутри ячейки, без отступа
    первой строки, выровнен по центру внутри ячейки.
    """
    lines = [ln.rstrip() for ln in approval_text.splitlines() if ln.strip()]
    if not lines:
        return

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    table.columns[0].width = Cm(8.5)

    # Снимаем границы.
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    cell = table.cell(0, 0)
    cell.width = Cm(8.5)

    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        # Шапка-«УТВЕРЖДАЮ» (первая строка) обычно идёт жирным.
        run.bold = (i == 0)


def _render_title_body(doc, body_text: str):
    """
    Рендерит тело титульного листа с minimal markdown:
      • строка «# заголовок» → центрированный жирный заголовок (без отступа красной строки);
      • прочие непустые строки → обычные justified-абзацы с поддержкой **bold**;
      • пустые строки — пропускаются (разделение абзацев и так получается естественно).
    """
    if not body_text or not body_text.strip():
        return

    for raw_line in body_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.lstrip().startswith("# "):
            heading = line.lstrip()[2:].strip()
            _p_inline_bold(
                doc, f"**{heading}**",
                size=14,
                space_before=18, space_after=18,
                center=True, first_line_indent=False,
            )
        else:
            _p_inline_bold(doc, line, size=14, space_after=4)


def _render_title_page(doc) -> bool:
    """
    Рисует титульный лист в начале документа на основе настроек:
      • cfg.title_page_approval — гриф «УТВЕРЖДАЮ» (плашка справа);
      • cfg.title_page_body     — заголовок отчёта + вводный текст + реквизиты.
    Возвращает True, если что-то было нарисовано (тогда вызывающий код
    добавит page-break перед основным содержимым).
    Если оба поля пустые — титульный лист не создаётся.
    """
    approval = (cfg.get("title_page_approval") or "").strip()
    body = (cfg.get("title_page_body") or "").strip()
    if not approval and not body:
        return False

    if approval:
        _render_approval_stamp(doc, approval)
        # Небольшой воздушный отступ после плашки перед заголовком отчёта.
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

    if body:
        _render_title_body(doc, body)

    return True


# ===================== ДАННЫЕ (файл 1) =====================

def generate_data_docx(questions: list) -> bytes:
    """Файл со списками вопросов и статистикой без аналитики."""
    doc = _make_doc()
    _last_section = None

    for q in questions:
        sec = q.get("section")
        sec_name = sec.get("name") if sec else None

        if sec_name and sec_name != _last_section:
            is_first_section = (_last_section is None)
            _last_section = sec_name
            if not is_first_section:
                doc.add_page_break()
            _p(doc, sec_name, bold=True, size=14,
               space_before=(0 if is_first_section else 14),
               space_after=14, center=True, first_line_indent=False)
            if sec.get("description"):
                _p(doc, sec["description"], size=14, space_after=6)

        file_keys = q["file_keys"]
        file_labels = q["file_labels"]
        file_totals = q["file_totals"]
        is_single = len(file_keys) == 1

        for row in q["rows"]:
            if is_single:
                fk = file_keys[0]
                count = row["counts"].get(fk, 0)
                total = file_totals.get(fk, 0)
                pct = f"{count / total * 100:.1f}%" if total > 0 else "—"
                _p(doc, f"  • {row['answer']}: {count} ({pct})", space_after=1)
            else:
                parts = []
                for fk in file_keys:
                    label = file_labels.get(fk, fk)
                    count = row["counts"].get(fk, 0)
                    total = file_totals.get(fk, 0)
                    pct = f"{count / total * 100:.1f}%" if total > 0 else "—"
                    parts.append(f"{label}: {count} ({pct})")
                _p(doc, f"  • {row['answer']}: {'; '.join(parts)}", space_after=1)

        if q.get("show_total", True):
            if is_single:
                fk = file_keys[0]
                _p(doc, f"  Всего: {file_totals.get(fk, 0)}", bold=True, space_after=6)
            else:
                parts = [
                    f"{file_labels.get(fk, fk)}: {file_totals.get(fk, 0)}"
                    for fk in file_keys
                ]
                _p(doc, f"  Всего: {'; '.join(parts)}", bold=True, space_after=6)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ===================== ПРОМПТЫ =====================

def _build_question_prompt(
    question: dict,
    sec_name: str = "",
    sec_description: str = "",
) -> tuple[str, str]:

    # ── SYSTEM ──────────────────────────────────────────────────────────
    system_parts = cfg.get("prompt_system_role").splitlines()

    if sec_description:
        system_parts += [
            "",
            "## Контекст текущего раздела",
            sec_description,
            "",
            "Используй этот контекст при написании анализа:",
            "— Если здесь сформулировано конкретное требование или акцент — строго следуй ему.",
            "— Если это описание тематики или состава раздела — учитывай при интерпретации.",
            "— Если это пояснение или замечание — прими во внимание как фоновый контекст.",
            "Академический стиль — инструмент, он не должен вступать в противоречие с контекстом раздела.",
        ]

    system_prompt = "\n".join(system_parts)

    # ── USER ─────────────────────────────────────────────────────────────
    lines = [
        "Ниже — примеры желаемого стиля:",
        cfg.get("prompt_style_example"),
        "",
        "ПРАВИЛА НАПИСАНИЯ:",
        cfg.get("prompt_writing_rules"),
        "",
    ]

    if sec_name:
        lines += [f"Раздел анкеты: «{sec_name}»", ""]

    q = question
    lines += [
        f"Напиши аналитический фрагмент по вопросу: «{q['question_name']}»",
        "",
        "Статистика ответов:",
    ]
    for row in q["rows"]:
        parts = []
        for fk in q["file_keys"]:
            label = q["file_labels"].get(fk, fk)
            count = row["counts"].get(fk, 0)
            total = q["file_totals"].get(fk, 0)
            pct = round(count / total * 100, 1) if total > 0 else 0
            parts.append(f"{label}: {count} ({pct}%)")
        lines.append(f"  - {row['answer']}: {'; '.join(parts)}")

    if sec_description:
        lines += [
            "",
            "При написании учти контекст раздела, указанный в начале.",
        ]

    user_prompt = "\n".join(lines)
    return system_prompt, user_prompt


# ===================== ВЫЗОВ МОДЕЛИ =====================

def _call_llm(prompt: str, system: str = "", max_tokens: int | None = None) -> str:
    client = get_openrouter_client()
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    if max_tokens is None:
        max_tokens = cfg.get_int("llm_max_tokens_analysis")

    _pace()
    for attempt in range(6):
        try:
            response = client.chat.completions.create(
                model=cfg.get("openrouter_model"),
                messages=messages,
                max_tokens=max_tokens,
                temperature=cfg.get_float("llm_temperature"),
            )
            _mark_request_done()
            return response.choices[0].message.content.strip()
        except Exception as e:
            if "429" in str(e) and attempt < 5:
                _backoff_wait(attempt)
            else:
                raise


# ===================== ПАРАЛЛЕЛЬНАЯ ГЕНЕРАЦИЯ ТЕКСТОВ =====================

# Максимум одновременных запросов к OpenRouter. Для free-tier 3 — безопасный
# потолок: больше — будет много 429.
_MAX_CONCURRENCY = 3


def _generate_texts_parallel(
    questions: list,
    cancel_event=None,
    progress_callback=None,
    progress_total: int | None = None,
) -> list[dict]:
    """
    Параллельно генерирует аналитические тексты для всех вопросов.
    Возвращает список словарей в том же порядке, что и questions:
      {'text': str | None, 'skipped': bool, 'error': bool}
    """
    semaphore = threading.Semaphore(_MAX_CONCURRENCY)
    results: list[dict | None] = [None] * len(questions)
    done_counter = {"n": 0}
    counter_lock = threading.Lock()
    total = progress_total if progress_total is not None else len(questions)

    def _worker(idx: int, q: dict):
        if cancel_event and cancel_event.is_set():
            results[idx] = {"text": None, "skipped": True, "error": False}
            return

        # Защитный no-op: если когда-нибудь добавим skip_analytics — будет работать.
        if q.get("skip_analytics", False):
            results[idx] = {"text": None, "skipped": True, "error": False}
            return

        sec = q.get("section")
        sec_name = sec.get("name") if sec else ""
        sec_description = sec.get("description", "") if sec else ""

        # Проверка кэша до захвата семафора — кэш-хиты не блокируют слоты.
        key = _cache_key(q, sec_name, sec_description)
        with _analysis_cache_lock:
            cached = _analysis_cache.get(key)
        if cached is not None:
            results[idx] = {"text": cached, "skipped": False, "error": False}
            print(f"  [{idx + 1}/{total}] CACHE: {q.get('question_name', '')[:40]}")
        else:
            with semaphore:
                if cancel_event and cancel_event.is_set():
                    results[idx] = {"text": None, "skipped": True, "error": False}
                    return
                try:
                    system_prompt, user_prompt = _build_question_prompt(
                        q, sec_name, sec_description
                    )
                    text = _call_llm(user_prompt, system=system_prompt)
                    with _analysis_cache_lock:
                        _analysis_cache[key] = text
                    results[idx] = {"text": text, "skipped": False, "error": False}
                    print(f"  [{idx + 1}/{total}] OK: {q.get('question_name', '')[:40]}")
                except Exception as e:
                    results[idx] = {
                        "text": f"Ошибка генерации аналитики: {e}",
                        "skipped": False,
                        "error": True,
                    }
                    print(f"  [{idx + 1}/{total}] ERROR: {e}")

        # Обновляем прогресс по мере фактического завершения (а не по порядку).
        with counter_lock:
            done_counter["n"] += 1
            done_now = done_counter["n"]
        if progress_callback:
            try:
                progress_callback(done_now, total, q.get("question_name", ""))
            except Exception:
                pass

    threads: list[threading.Thread] = []
    for i, q in enumerate(questions):
        t = threading.Thread(target=_worker, args=(i, q), daemon=True)
        threads.append(t)
        t.start()

    for t in threads:
        t.join()

    # На всякий случай: если какой-то слот остался None (теоретически невозможно).
    for i, r in enumerate(results):
        if r is None:
            results[i] = {"text": None, "skipped": True, "error": False}
    return results  # type: ignore[return-value]


# ===================== ВЫВОДЫ ПО РАЗДЕЛАМ =====================

_section_conclusion_cache: dict[str, str] = {}
_section_conclusion_cache_lock = threading.Lock()


def _section_conclusion_cache_key(sec_name: str, sec_desc: str, qs: list[dict]) -> str:
    payload = json.dumps(
        {
            "name":  sec_name,
            "desc":  sec_desc,
            "qs":    [
                {"q": q.get("question_name"), "rows": q.get("rows")}
                for q in qs
            ],
            "model": cfg.get("openrouter_model"),
            "sys":   hashlib.md5(cfg.get("prompt_section_conclusion_system").encode()).hexdigest(),
            "rules": hashlib.md5(cfg.get("prompt_section_conclusion_rules").encode()).hexdigest(),
            "ex":    hashlib.md5(cfg.get("prompt_section_conclusion_example").encode()).hexdigest(),
        },
        ensure_ascii=False,
        sort_keys=True,
        default=str,
    )
    return hashlib.md5(payload.encode("utf-8")).hexdigest()


def _build_section_conclusion_prompt(
    sec_name: str,
    sec_description: str,
    questions_in_section: list[dict],
) -> tuple[str, str]:
    """Собирает system + user-промпт для выводов по разделу."""

    # ── SYSTEM ──────────────────────────────────────────────────────────
    system_parts = cfg.get("prompt_section_conclusion_system").splitlines()
    if sec_description:
        system_parts += [
            "",
            "## Контекст текущего раздела",
            sec_description,
            "",
            "Используй этот контекст при формулировании выводов:",
            "— Если здесь сформулировано конкретное требование или акцент — строго следуй ему "
            "при выборе того, что вынести в выводы.",
            "— Если это описание тематики или состава раздела — учитывай при интерпретации "
            "статистики и формулировке обобщений.",
            "— Если это пояснение или замечание — прими во внимание как фоновый контекст.",
            "Академический стиль — инструмент, он не должен вступать в противоречие с контекстом раздела.",
        ]
    system_prompt = "\n".join(system_parts)

    # ── USER ─────────────────────────────────────────────────────────────
    lines = [
        "Ниже — пример желаемого стиля выводов:",
        cfg.get("prompt_section_conclusion_example"),
        "",
        "ПРАВИЛА НАПИСАНИЯ ВЫВОДОВ:",
        cfg.get("prompt_section_conclusion_rules"),
        "",
        f"Раздел отчёта: «{sec_name}»",
        "",
    ]

    if sec_description:
        lines += [
            "Описание раздела:",
            sec_description,
            "",
        ]

    lines += [
        "Ниже — все вопросы этого раздела со статистикой ответов. "
        "Сформулируй на основании СОВОКУПНОСТИ этих данных обобщающие выводы по разделу"
        + (", опираясь в том числе на описание раздела выше." if sec_description else "."),
        "",
    ]

    for q in questions_in_section:
        lines.append(f"— Вопрос: «{q.get('question_name', '')}»")
        for row in q.get("rows", []):
            parts = []
            for fk in q.get("file_keys", []):
                label = q.get("file_labels", {}).get(fk, fk)
                count = row["counts"].get(fk, 0)
                total = q.get("file_totals", {}).get(fk, 0)
                pct = round(count / total * 100, 1) if total > 0 else 0
                parts.append(f"{label}: {count} ({pct}%)")
            lines.append(f"    {row.get('answer', '')}: {'; '.join(parts)}")
        lines.append("")

    if sec_description:
        lines.append(
            "Перед формулировкой выводов сверься с описанием раздела выше — "
            "оно задаёт акцент и рамку интерпретации данных."
        )
    lines.append(
        "Напиши только связный текст выводов — без заголовка, без маркеров списков, "
        "без нумерации, без преамбулы вроде «Итак,» или «В заключение,»."
    )

    return system_prompt, "\n".join(lines)


def _generate_section_conclusions_parallel(
    sections: list[dict],
    cancel_event=None,
    progress_callback=None,
    progress_offset: int = 0,
    progress_total: int | None = None,
) -> list[dict]:
    """
    sections — упорядоченный список вида:
       [{"name": str, "description": str, "questions": [q, q, ...]}, ...]
    Возвращает параллельный список:
       [{"text": str | None, "skipped": bool, "error": bool}, ...]
    """
    semaphore = threading.Semaphore(_MAX_CONCURRENCY)
    results: list[dict | None] = [None] * len(sections)
    done_counter = {"n": progress_offset}
    counter_lock = threading.Lock()
    total_for_progress = progress_total if progress_total is not None else len(sections)
    max_tokens = cfg.get_int("llm_max_tokens_section_conclusion")

    def _worker(idx: int, sec: dict):
        if cancel_event and cancel_event.is_set():
            results[idx] = {"text": None, "skipped": True, "error": False}
            return

        name = sec.get("name") or ""
        desc = sec.get("description") or ""
        qs = sec.get("questions") or []

        key = _section_conclusion_cache_key(name, desc, qs)
        with _section_conclusion_cache_lock:
            cached = _section_conclusion_cache.get(key)
        if cached is not None:
            results[idx] = {"text": cached, "skipped": False, "error": False}
            print(f"  [SEC CACHE] Выводы раздела «{name[:40]}»")
        else:
            with semaphore:
                if cancel_event and cancel_event.is_set():
                    results[idx] = {"text": None, "skipped": True, "error": False}
                    return
                try:
                    sys_p, usr_p = _build_section_conclusion_prompt(name, desc, qs)
                    text = _call_llm(usr_p, system=sys_p, max_tokens=max_tokens)
                    with _section_conclusion_cache_lock:
                        _section_conclusion_cache[key] = text
                    results[idx] = {"text": text, "skipped": False, "error": False}
                    print(f"  [SEC OK] Выводы раздела «{name[:40]}»")
                except Exception as e:
                    results[idx] = {
                        "text": f"Ошибка генерации выводов: {e}",
                        "skipped": False,
                        "error": True,
                    }
                    print(f"  [SEC ERROR] «{name[:40]}»: {e}")

        with counter_lock:
            done_counter["n"] += 1
            done_now = done_counter["n"]
        if progress_callback:
            try:
                progress_callback(done_now, total_for_progress, f"Выводы: «{name}»")
            except Exception:
                pass

    threads: list[threading.Thread] = []
    for i, sec in enumerate(sections):
        t = threading.Thread(target=_worker, args=(i, sec), daemon=True)
        threads.append(t)
        t.start()

    for t in threads:
        t.join()

    for i, r in enumerate(results):
        if r is None:
            results[i] = {"text": None, "skipped": True, "error": False}
    return results  # type: ignore[return-value]


# ===================== ИТОГОВЫЕ ВЫВОДЫ ПО ОПРОСУ =====================

_final_conclusion_cache: dict[str, str] = {}
_final_conclusion_cache_lock = threading.Lock()


def _final_conclusion_cache_key(questions: list[dict]) -> str:
    payload = json.dumps(
        {
            "qs": [
                {
                    "q":    q.get("question_name"),
                    "rows": q.get("rows"),
                    "sec":  (q.get("section") or {}).get("name") if q.get("section") else None,
                }
                for q in questions
            ],
            "model": cfg.get("openrouter_model"),
            "sys":   hashlib.md5(cfg.get("prompt_final_conclusion_system").encode()).hexdigest(),
            "rules": hashlib.md5(cfg.get("prompt_final_conclusion_rules").encode()).hexdigest(),
            "ex":    hashlib.md5(cfg.get("prompt_final_conclusion_example").encode()).hexdigest(),
        },
        ensure_ascii=False,
        sort_keys=True,
        default=str,
    )
    return hashlib.md5(payload.encode("utf-8")).hexdigest()


def _build_final_conclusion_prompt(questions: list[dict]) -> tuple[str, str]:
    """
    Собирает system + user-промпт для итоговых выводов по всему опросу.
    В user-блок попадают все вопросы со статистикой ответов, сгруппированные по разделам,
    в исходном порядке появления.
    """
    system_prompt = cfg.get("prompt_final_conclusion_system")

    lines = [
        "Ниже — пример желаемого стиля одного из итоговых выводов:",
        cfg.get("prompt_final_conclusion_example"),
        "",
        "ПРАВИЛА НАПИСАНИЯ ИТОГОВЫХ ВЫВОДОВ:",
        cfg.get("prompt_final_conclusion_rules"),
        "",
        "Ниже — все вопросы опроса со статистикой ответов, "
        "сгруппированные по разделам. Сформулируй на их основе итоговые выводы.",
        "",
    ]

    # Группировка вопросов по разделам — для лучшей читаемости моделью.
    current_sec_name: str | None | object = "__SENTINEL__"
    for q in questions:
        sec = q.get("section")
        sec_name = sec.get("name") if sec else None

        if sec_name != current_sec_name:
            current_sec_name = sec_name
            lines.append("")
            if sec_name:
                lines.append(f"## Раздел: «{sec_name}»")
                desc = (sec or {}).get("description") or ""
                if desc:
                    lines.append(f"Описание раздела: {desc}")
            else:
                lines.append("## Вопросы вне разделов")
            lines.append("")

        lines.append(f"— Вопрос: «{q.get('question_name', '')}»")
        for row in q.get("rows", []):
            parts = []
            for fk in q.get("file_keys", []):
                label = q.get("file_labels", {}).get(fk, fk)
                count = row["counts"].get(fk, 0)
                total = q.get("file_totals", {}).get(fk, 0)
                pct = round(count / total * 100, 1) if total > 0 else 0
                parts.append(f"{label}: {count} ({pct}%)")
            lines.append(f"    {row.get('answer', '')}: {'; '.join(parts)}")
        lines.append("")

    lines += [
        "Напиши итоговые выводы: РОВНО 4–5 связных абзацев по 5–6 предложений каждый, "
        "разделённых пустой строкой. Без общего заголовка, без нумерации выводов, "
        "без маркеров списка. Обязательно укажи, что респонденты предлагают ВВЕСТИ "
        "или развивать в работе университета и от чего предлагают ОТКАЗАТЬСЯ.",
    ]

    return system_prompt, "\n".join(lines)


def _generate_final_conclusion(
    questions: list[dict],
    cancel_event=None,
    progress_callback=None,
    progress_offset: int = 0,
    progress_total: int | None = None,
) -> dict:
    """Один LLM-вызов на финальные выводы. Кэшируется в памяти процесса."""
    if cancel_event and cancel_event.is_set():
        return {"text": None, "skipped": True, "error": False}

    key = _final_conclusion_cache_key(questions)
    with _final_conclusion_cache_lock:
        cached = _final_conclusion_cache.get(key)

    if cached is not None:
        print("  [FINAL CACHE] Итоговые выводы")
        result = {"text": cached, "skipped": False, "error": False}
    else:
        try:
            sys_p, usr_p = _build_final_conclusion_prompt(questions)
            text = _call_llm(
                usr_p,
                system=sys_p,
                max_tokens=cfg.get_int("llm_max_tokens_final_conclusion"),
            )
            with _final_conclusion_cache_lock:
                _final_conclusion_cache[key] = text
            result = {"text": text, "skipped": False, "error": False}
            print("  [FINAL OK] Итоговые выводы")
        except Exception as e:
            result = {
                "text": f"Ошибка генерации итоговых выводов: {e}",
                "skipped": False,
                "error": True,
            }
            print(f"  [FINAL ERROR] {e}")

    if progress_callback and progress_total:
        try:
            progress_callback(progress_offset + 1, progress_total, "Итоговые выводы")
        except Exception:
            pass

    return result


# ===================== АНАЛИТИКА (файл 2) =====================

def _group_questions_by_section(questions: list) -> list[tuple[dict | None, list[dict]]]:
    """
    Группирует вопросы в исходном порядке: при смене section.name открывается
    новая группа. Возвращает список (section_dict_or_None, [questions_in_group]).
    Вопросы без раздела попадают в группу с section=None.
    """
    groups: list[tuple[dict | None, list[dict]]] = []
    current_key: object = "__SENTINEL__"
    for q in questions:
        sec = q.get("section")
        key = sec.get("name") if sec else None
        if key != current_key:
            groups.append((sec, []))
            current_key = key
        groups[-1][1].append(q)
    return groups


def generate_analysis_docx(questions: list, progress_callback=None, cancel_event=None) -> bytes:
    """
    Аналитический файл. LLM-вызовы выполняются параллельно (до _MAX_CONCURRENCY),
    после чего документ собирается строго в исходном порядке вопросов.
    После всех вопросов раздела вставляется блок «Выводы раздела «...»» —
    отдельный LLM-вызов на основании совокупности вопросов раздела.
    Для вопросов без раздела выводы не формируются.
    """
    doc = _make_doc()
    total_questions = len(questions)
    chart_counter = [1]
    table_counter = [1]
    part_counter = [1]

    # 0) Титульный лист — если хотя бы одно из полей в настройках заполнено.
    title_rendered = _render_title_page(doc)
    if title_rendered:
        doc.add_page_break()

    # 1) Группировка вопросов по разделам в порядке появления.
    section_groups = _group_questions_by_section(questions)

    # 2) Отбираем разделы, для которых нужны выводы (есть имя раздела и
    #    хотя бы один вопрос). Сохраняем порядок появления.
    sections_for_conclusion = [
        {
            "name":        sec.get("name", ""),
            "description": sec.get("description", "") if sec else "",
            "questions":   qs,
        }
        for sec, qs in section_groups
        if sec and sec.get("name") and qs
    ]
    total_conclusions = len(sections_for_conclusion)
    total_final = 1 if questions else 0  # один LLM-вызов на финальные выводы

    print(
        f"Параллельная генерация {total_questions} вопросов, "
        f"{total_conclusions} выводов по разделам и "
        f"{total_final} итогового вывода "
        f"(до {_MAX_CONCURRENCY} потоков)..."
    )

    grand_total = total_questions + total_conclusions + total_final

    # 3) Параллельно генерируем все тексты вопросов.
    texts = _generate_texts_parallel(
        questions,
        cancel_event=cancel_event,
        progress_callback=progress_callback,
        # шкала прогресса единая на весь экспорт: N вопросов + M выводов
        progress_total=grand_total,
    )

    # 4) Параллельно: выводы по разделам и финальные выводы. Оба этапа не
    #    зависят от текстов вопросов и не зависят друг от друга — запускаем
    #    одновременно в двух потоках. Внутренние пулы делят rate-лимит через
    #    общий _pace() / _LAST_REQUEST_TIME, так что 429 нам не грозят.
    conclusions: list[dict] = []
    final_result: dict | None = None

    def _run_section_conclusions():
        nonlocal conclusions
        if sections_for_conclusion and not (cancel_event and cancel_event.is_set()):
            conclusions = _generate_section_conclusions_parallel(
                sections_for_conclusion,
                cancel_event=cancel_event,
                progress_callback=progress_callback,
                progress_offset=total_questions,
                progress_total=grand_total,
            )

    def _run_final_conclusion():
        nonlocal final_result
        if total_final and not (cancel_event and cancel_event.is_set()):
            final_result = _generate_final_conclusion(
                questions,
                cancel_event=cancel_event,
                progress_callback=progress_callback,
                progress_offset=total_questions + total_conclusions,
                progress_total=grand_total,
            )

    t_sec = threading.Thread(target=_run_section_conclusions, daemon=True)
    t_fin = threading.Thread(target=_run_final_conclusion, daemon=True)
    t_sec.start()
    t_fin.start()
    t_sec.join()
    t_fin.join()

    # Сопоставляем сгенерированные выводы их разделам через имя раздела
    # (имена в пределах отчёта считаем уникальными, как и в текущей сборке).
    conclusion_by_section_name: dict[str, dict] = {}
    for sec_def, res in zip(sections_for_conclusion, conclusions):
        conclusion_by_section_name[sec_def["name"]] = res

    # 5) Сборка документа: проходим по группам в исходном порядке.
    is_first_section_written = True
    for group_idx, (sec, qs_in_group) in enumerate(section_groups):
        if cancel_event and cancel_event.is_set():
            break

        sec_name = sec.get("name") if sec else ""
        sec_description = sec.get("description", "") if sec else ""

        # Заголовок раздела (только для именованных разделов).
        if sec_name:
            if not is_first_section_written:
                doc.add_page_break()
            _p(
                doc, sec_name, bold=True, size=14,
                space_before=(0 if is_first_section_written else 14),
                space_after=14, center=True, first_line_indent=False,
            )
            if sec_description:
                _p(doc, sec_description, size=14, space_after=6)
            is_first_section_written = False

        # Вопросы раздела.
        for q in qs_in_group:
            if cancel_event and cancel_event.is_set():
                break
            # индекс вопроса в исходном списке — для попадания в texts
            q_idx = questions.index(q)
            result = texts[q_idx] or {}
            if not result.get("skipped") and result.get("text"):
                bold = bool(result.get("error"))
                for para_text in [p.strip() for p in result["text"].split("\n\n") if p.strip()]:
                    _p(doc, para_text, bold=bold, size=14, space_after=6)

            insert_visualization(doc, q, chart_counter, table_counter, part_counter)

        # Выводы раздела — только для именованных разделов.
        if sec_name and sec_name in conclusion_by_section_name:
            conc = conclusion_by_section_name[sec_name]
            if conc and not conc.get("skipped") and conc.get("text"):
                _p(
                    doc, f"Выводы раздела «{sec_name}»",
                    bold=True, size=14,
                    space_before=12, space_after=8,
                    first_line_indent=False,
                )
                bold_text = bool(conc.get("error"))
                for para_text in [p.strip() for p in conc["text"].split("\n\n") if p.strip()]:
                    _p(doc, para_text, bold=bold_text, size=14, space_after=6)

    # 6) ИТОГОВЫЕ ВЫВОДЫ — отдельная последняя страница.
    if (
        final_result
        and not final_result.get("skipped")
        and final_result.get("text")
        and not (cancel_event and cancel_event.is_set())
    ):
        doc.add_page_break()
        _p(
            doc, "Итоговые выводы по результатам опроса",
            bold=True, size=14,
            space_before=0, space_after=14,
            center=True, first_line_indent=False,
        )
        bold_err = bool(final_result.get("error"))
        for para_text in [p.strip() for p in final_result["text"].split("\n\n") if p.strip()]:
            _p(doc, para_text, bold=bold_err, size=14, space_after=8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ===================== ЕДИНАЯ ТОЧКА ВХОДА =====================

def generate_docx(questions: list, progress_callback=None) -> tuple[bytes, bytes]:
    data_bytes = generate_data_docx(questions)
    analysis_bytes = generate_analysis_docx(questions, progress_callback=progress_callback)
    return data_bytes, analysis_bytes
