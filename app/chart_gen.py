# ===================== chart_gen.py =====================
"""
Embedding OOXML charts and Word tables into python-docx documents.
Charts are embedded as proper Excel charts (editable by double-click in Word),
not as images.
"""
import io
import openpyxl
from lxml import etree
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

RT_CHART   = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart'
RT_PACKAGE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/package'
CT_CHART   = 'application/vnd.openxmlformats-officedocument.drawingml.chart+xml'
CT_EXCEL   = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

# 1 cm = 360 000 EMU. Text width = 21 - 3 - 1.5 = 16.5 cm
TEXT_W = 5_940_000   # 16.5 cm
BAR_H  = 3_600_000   # 10 cm
PIE_H  = 3_240_000   #  9 cm


# ── helpers ──────────────────────────────────────────────────────────────────

def _x(s: object) -> str:
    return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def _color_hex(c) -> str:
    """Привести цвет к OOXML-формату: 6 hex-символов без '#'."""
    if not c or not isinstance(c, str):
        return ''
    c = c.strip().lstrip('#')
    if len(c) == 3:
        c = ''.join(ch * 2 for ch in c)
    if len(c) < 6:
        return ''
    try:
        int(c[:6], 16)
    except ValueError:
        return ''
    return c[:6].upper()


def _data_labels_block(pos: str, show_val: bool, show_percent: bool,
                       num_fmt: str = '0&quot;%&quot;',
                       size: int = 1200, color_hex: str = '000000') -> list:
    """Возвращает строки <c:dLbls> для серии: жирные подписи с процентом."""
    return [
        '          <c:dLbls>',
        f'            <c:numFmt formatCode="{num_fmt}" sourceLinked="0"/>',
        '            <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>',
        '            <c:txPr>',
        '              <a:bodyPr wrap="square" lIns="38100" tIns="19050" rIns="38100" bIns="19050" anchor="ctr"/>',
        '              <a:lstStyle/>',
        '              <a:p>',
        f'                <a:pPr><a:defRPr sz="{size}" b="1">'
        f'<a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill>'
        f'<a:latin typeface="Times New Roman"/></a:defRPr></a:pPr>',
        '                <a:endParaRPr lang="ru-RU"/>',
        '              </a:p>',
        '            </c:txPr>',
        f'            <c:dLblPos val="{pos}"/>',
        '            <c:showLegendKey val="0"/>',
        f'            <c:showVal val="{1 if show_val else 0}"/>',
        '            <c:showCatName val="0"/>',
        '            <c:showSerName val="0"/>',
        f'            <c:showPercent val="{1 if show_percent else 0}"/>',
        '            <c:showBubbleSize val="0"/>',
        '          </c:dLbls>',
    ]


def _values_to_percent_per_series(series_values):
    """Конвертирует [[counts...], ...] → [[pct...], ...] по сумме своей серии (как на экране)."""
    out = []
    for vals in series_values:
        total = sum(vals)
        if total > 0:
            out.append([round(v / total * 100, 2) for v in vals])
        else:
            out.append([0 for _ in vals])
    return out


def _build_xlsx(answers: list, series_labels: list, series_values: list) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Sheet1'
    ws.cell(row=1, column=1, value='')
    for ci, lbl in enumerate(series_labels, 2):
        ws.cell(row=1, column=ci, value=lbl)
    for ri, ans in enumerate(answers):
        ws.cell(row=ri + 2, column=1, value=ans)
        for ci, vals in enumerate(series_values):
            ws.cell(row=ri + 2, column=ci + 2, value=vals[ri])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── chart XML builders ───────────────────────────────────────────────────────

def _bar_xml(answers, series_labels, series_values,
             bar_dir='col', stacked=False, show_legend=True,
             series_colors=None, point_colors=None) -> bytes:
    """
    series_colors: список hex-цветов по сериям (multi-file) — раскрашивает каждую серию своим цветом.
    point_colors:  список hex-цветов по точкам (single-file) — раскрашивает каждый столбик своим цветом.
    stacked: True → percentStacked (все столбцы 100%, как в drawStackedChart).
    """
    n = len(answers)
    grouping = 'percentStacked' if stacked else 'clustered'
    cat_pos, val_pos = ('l', 'b') if bar_dir == 'bar' else ('b', 'l')

    out = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        '  <c:lang val="ru-RU"/>',
        '  <c:chart>',
        '    <c:autoTitleDeleted val="1"/>',
        '    <c:plotArea><c:layout/>',
        '      <c:barChart>',
        f'        <c:barDir val="{bar_dir}"/>',
        f'        <c:grouping val="{grouping}"/>',
        '        <c:varyColors val="0"/>',
    ]

    for si, (lbl, vals) in enumerate(zip(series_labels, series_values)):
        col = chr(ord('B') + si)

        out += [
            f'        <c:ser>',
            f'          <c:idx val="{si}"/><c:order val="{si}"/>',
            f'          <c:tx><c:strRef><c:f>Sheet1!${col}$1</c:f>',
            f'            <c:strCache><c:ptCount val="1"/>',
            f'              <c:pt idx="0"><c:v>{_x(lbl)}</c:v></c:pt>',
            f'            </c:strCache></c:strRef></c:tx>',
        ]

        series_hex = ''
        if series_colors and si < len(series_colors):
            series_hex = _color_hex(series_colors[si])
        if series_hex:
            out.append(
                f'          <c:spPr><a:solidFill><a:srgbClr val="{series_hex}"/></a:solidFill>'
                f'<a:ln><a:noFill/></a:ln></c:spPr>'
            )
        out.append('          <c:invertIfNegative val="0"/>')

        if point_colors and not stacked:
            for pi, pc in enumerate(point_colors):
                phex = _color_hex(pc)
                if not phex:
                    continue
                out += [
                    f'          <c:dPt>',
                    f'            <c:idx val="{pi}"/>',
                    f'            <c:invertIfNegative val="0"/>',
                    f'            <c:bubble3D val="0"/>',
                    f'            <c:spPr><a:solidFill><a:srgbClr val="{phex}"/></a:solidFill>'
                    f'<a:ln><a:noFill/></a:ln></c:spPr>',
                    f'          </c:dPt>',
                ]

        out += _data_labels_block(
            pos='ctr' if stacked else 'outEnd',
            show_val=True,
            show_percent=False,
        )
        # Категории
        out += [
            f'          <c:cat><c:strRef>',
            f'            <c:f>Sheet1!$A$2:$A${n+1}</c:f>',
            f'            <c:strCache><c:ptCount val="{n}"/>',
        ]
        for i, ans in enumerate(answers):
            out.append(f'              <c:pt idx="{i}"><c:v>{_x(ans)}</c:v></c:pt>')
        out += [
            f'            </c:strCache></c:strRef></c:cat>',
            f'          <c:val><c:numRef>',
            f'            <c:f>Sheet1!${col}$2:${col}${n+1}</c:f>',
            f'            <c:numCache><c:formatCode>General</c:formatCode>',
            f'              <c:ptCount val="{n}"/>',
        ]
        for i, v in enumerate(vals):
            out.append(f'              <c:pt idx="{i}"><c:v>{v}</c:v></c:pt>')
        out += [
            f'            </c:numCache></c:numRef></c:val>',
            f'        </c:ser>',
        ]


    if stacked:
        out += [
            '        <c:gapWidth val="60"/>',
            '        <c:overlap val="100"/>',
        ]
    else:
        out.append('        <c:gapWidth val="80"/>')

    out += [
        '        <c:axId val="100"/><c:axId val="101"/>',
        '      </c:barChart>',
        '      <c:catAx>',
        '        <c:axId val="100"/>',
        '        <c:scaling><c:orientation val="minMax"/></c:scaling>',
        '        <c:delete val="0"/>',
        f'        <c:axPos val="{cat_pos}"/>',
        '        <c:majorTickMark val="none"/>',
        '        <c:minorTickMark val="none"/>',
        '        <c:crossAx val="101"/>',
        '      </c:catAx>',
        '      <c:valAx>',
        '        <c:axId val="101"/>',
        '        <c:scaling><c:orientation val="minMax"/></c:scaling>',
        '        <c:delete val="1"/>',
        f'        <c:axPos val="{val_pos}"/>',
        '        <c:crossAx val="100"/>',
        '      </c:valAx>',
        '    </c:plotArea>',
    ]

    effective_show_legend = show_legend and not (point_colors and not stacked)
    if effective_show_legend:

        out += [
            '    <c:legend>',
            '      <c:legendPos val="b"/>',
            '      <c:layout>',
            '        <c:manualLayout>',
            '          <c:xMode val="edge"/><c:yMode val="edge"/>',
            '          <c:x val="0"/><c:y val="0.88"/>',
            '          <c:w val="1"/><c:h val="0.12"/>',
            '        </c:manualLayout>',
            '      </c:layout>',
            '      <c:overlay val="0"/>',
            '    </c:legend>',
        ]
    out += [
        '    <c:plotVisOnly val="1"/>',
        '  </c:chart>',
        '  <c:externalData r:id="rId1"><c:autoUpdate val="0"/></c:externalData>',
        '</c:chartSpace>',
    ]
    return '\n'.join(out).encode('utf-8')


def _pie_xml(answers, values, series_label, show_legend=True,
             point_colors=None) -> bytes:
    """
    point_colors: список hex-цветов для каждого сектора (как на экране в drawPieChart).
    """
    n = len(answers)
    out = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        '  <c:lang val="ru-RU"/>',
        '  <c:chart>',
        '    <c:autoTitleDeleted val="1"/>',
        '    <c:plotArea><c:layout/>',
        '      <c:pieChart>',
        '        <c:varyColors val="1"/>',
        '        <c:ser>',
        '          <c:idx val="0"/><c:order val="0"/>',
        '          <c:tx><c:strRef><c:f>Sheet1!$B$1</c:f>',
        '            <c:strCache><c:ptCount val="1"/>',
        f'              <c:pt idx="0"><c:v>{_x(series_label)}</c:v></c:pt>',
        '            </c:strCache></c:strRef></c:tx>',
    ]
    # Цвета секторов (OOXML: dPt идёт между tx/spPr и cat)
    if point_colors:
        for pi, pc in enumerate(point_colors):
            phex = _color_hex(pc)
            if not phex:
                continue
            out += [
                f'          <c:dPt>',
                f'            <c:idx val="{pi}"/>',
                f'            <c:bubble3D val="0"/>',
                f'            <c:spPr><a:solidFill><a:srgbClr val="{phex}"/></a:solidFill>'
                f'<a:ln w="9525"><a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill></a:ln></c:spPr>',
                f'          </c:dPt>',
            ]
    # Подписи: жирный чёрный процент. Для пирога OOXML сам вычислит % из сумм.
    out += _data_labels_block(
        pos='outEnd',
        show_val=False,
        show_percent=True,
        num_fmt='0%',
    )
    out += [
        '          <c:cat><c:strRef>',
        f'            <c:f>Sheet1!$A$2:$A${n+1}</c:f>',
        f'            <c:strCache><c:ptCount val="{n}"/>',
    ]
    for i, ans in enumerate(answers):
        out.append(f'              <c:pt idx="{i}"><c:v>{_x(ans)}</c:v></c:pt>')
    out += [
        '            </c:strCache></c:strRef></c:cat>',
        '          <c:val><c:numRef>',
        f'            <c:f>Sheet1!$B$2:$B${n+1}</c:f>',
        '            <c:numCache><c:formatCode>General</c:formatCode>',
        f'              <c:ptCount val="{n}"/>',
    ]
    for i, v in enumerate(values):
        out.append(f'              <c:pt idx="{i}"><c:v>{v}</c:v></c:pt>')
    out += [
        '            </c:numCache></c:numRef></c:val>',
        '        </c:ser>',
        '      </c:pieChart>',
        '    </c:plotArea>',
    ]
    if show_legend:
        out += [
            '    <c:legend>',
            '      <c:legendPos val="b"/>',
            '      <c:layout>',
            '        <c:manualLayout>',
            '          <c:xMode val="edge"/><c:yMode val="edge"/>',
            '          <c:x val="0"/><c:y val="0.85"/>',
            '          <c:w val="1"/><c:h val="0.15"/>',
            '        </c:manualLayout>',
            '      </c:layout>',
            '      <c:overlay val="0"/>',
            '    </c:legend>',
        ]
    out += [
        '    <c:plotVisOnly val="1"/>',
        '  </c:chart>',
        '  <c:externalData r:id="rId1"><c:autoUpdate val="0"/></c:externalData>',
        '</c:chartSpace>',
    ]
    return '\n'.join(out).encode('utf-8')


# ── OPC embedding ─────────────────────────────────────────────────────────────

def _embed_chart(doc, chart_xml: bytes, xlsx: bytes, n: int, cx=TEXT_W, cy=BAR_H):
    """Add a chart Part to the document and insert an inline drawing paragraph."""
    chart_part = Part(PackURI(f'/word/charts/chart{n}.xml'), CT_CHART, chart_xml)
    r_id = doc.part.relate_to(chart_part, RT_CHART)

    excel_part = Part(PackURI(f'/word/embeddings/sheet{n}.xlsx'), CT_EXCEL, xlsx)
    chart_part.relate_to(excel_part, RT_PACKAGE)   # always becomes rId1 in chart

    drawing_xml = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{n}" name="Chart {n}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">'
        f'<c:chart r:id="{r_id}"/>'
        f'</a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing>'
    )
    para = doc.add_paragraph()
    run = para.add_run()
    run._r.append(etree.fromstring(drawing_xml.encode('utf-8')))


# ── public API ────────────────────────────────────────────────────────────────

def insert_visualization(doc, q: dict, chart_counter: list,
                         table_counter: list = None, part_counter: list = None):
    """
    Insert the appropriate visualization (table or chart) after AI text.
    chart_counter — номер «Рисунок N» (один на вопрос).
    table_counter — номер «Таблица N» (один на таблицу).
    part_counter  — порядковый номер файла chart{N}.xml/sheet{N}.xlsx внутри docx
                    (всегда уникальный, иначе ZIP схватит дубликаты).
    """
    if table_counter is None:
        table_counter = [1]
    if part_counter is None:
        part_counter = [1]

    def _next_part():
        p = part_counter[0]
        part_counter[0] += 1
        return p

    viz_tab = q.get('viz_tab')
    if not viz_tab:
        return

    file_keys = q['file_keys']
    file_labels = q['file_labels']
    rows_data = q['rows']
    show_legend = q.get('show_legend', True)
    question_name = q.get('question_name', '')

    answers = [r['answer'] for r in rows_data]

    # На экране chartDirection='x' = вертикальные столбцы, 'y' = горизонтальные.
    # В OOXML: barDir='col' = вертикальные, 'bar' = горизонтальные.
    bar_dir = 'col' if q.get('chart_direction', 'y') == 'x' else 'bar'

    is_single_file = len(file_keys) == 1
    pie_colors = q.get('pie_colors') or []
    bar_colors = q.get('bar_colors') or []
    file_colors_map = q.get('file_colors') or {}
    series_colors = [file_colors_map.get(fk, '') for fk in file_keys]
    point_colors_bar = bar_colors if is_single_file else None

    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    def _add_caption(text, above=False, center=True):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2 if above else 4)
        para.paragraph_format.space_after = Pt(14)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return para

    fig_state = {'n': None}

    def _fig_num():
        if fig_state['n'] is None:
            fig_state['n'] = chart_counter[0]
            chart_counter[0] += 1
        return fig_state['n']

    if viz_tab == 'table':
        # Подпись НАД таблицей
        t_num = table_counter[0]
        table_counter[0] += 1
        _add_caption(
            f'Таблица {t_num} – Распределение ответов респондентов на вопрос: «{question_name}»',
            above=True
        )
        _insert_word_table(doc, q)

    elif viz_tab == 'both':
        # Подпись над таблицей
        t_num = table_counter[0]
        table_counter[0] += 1
        _add_caption(
            f'Таблица {t_num} – Распределение ответов респондентов на вопрос: «{question_name}»',
            above=True
        )
        _insert_word_table(doc, q)

        # Диаграмма того типа, что был выбран на фронте
        chart_type = q.get('both_chart_type', 'bar')
        series_labels = [file_labels.get(fk, fk) for fk in file_keys]
        series_counts = [[r['counts'].get(fk, 0) for r in rows_data] for fk in file_keys]
        series_values_pct = _values_to_percent_per_series(series_counts)

        if chart_type == 'pie':
            n = _fig_num()
            for fk in file_keys:
                values = [r['counts'].get(fk, 0) for r in rows_data]
                lbl = file_labels.get(fk, fk)
                chart_xml = _pie_xml(answers, values, lbl, show_legend,
                                     point_colors=pie_colors)
                xlsx = _build_xlsx(answers, [lbl], [values])
                _embed_chart(doc, chart_xml, xlsx, _next_part(), cy=PIE_H)
            _add_caption(
                f'Рисунок {n} – Распределение ответов респондентов на вопрос: «{question_name}»',
                above=False
            )
        else:
            stacked = (chart_type == 'stacked')
            chart_xml = _bar_xml(
                answers, series_labels, series_values_pct,
                bar_dir, stacked, show_legend,
                series_colors=series_colors,
                point_colors=point_colors_bar,
            )
            xlsx = _build_xlsx(answers, series_labels, series_values_pct)
            n = _fig_num()
            _embed_chart(doc, chart_xml, xlsx, _next_part())
            _add_caption(
                f'Рисунок {n} – Распределение ответов респондентов на вопрос: «{question_name}»',
                above=False
            )

    elif viz_tab in ('bar', 'stacked'):
        series_labels = [file_labels.get(fk, fk) for fk in file_keys]
        series_counts = [[r['counts'].get(fk, 0) for r in rows_data] for fk in file_keys]
        series_values_pct = _values_to_percent_per_series(series_counts)
        stacked = (viz_tab == 'stacked')
        chart_xml = _bar_xml(
            answers, series_labels, series_values_pct,
            bar_dir, stacked, show_legend,
            series_colors=series_colors,
            point_colors=point_colors_bar,
        )
        xlsx = _build_xlsx(answers, series_labels, series_values_pct)
        n = _fig_num()
        _embed_chart(doc, chart_xml, xlsx, _next_part())
        _add_caption(
            f'Рисунок {n} – Распределение ответов респондентов на вопрос: «{question_name}»',
            above=False
        )

    elif viz_tab == 'pie':
        n = _fig_num()
        for fk in file_keys:
            values = [r['counts'].get(fk, 0) for r in rows_data]
            lbl = file_labels.get(fk, fk)
            chart_xml = _pie_xml(answers, values, lbl, show_legend,
                                 point_colors=pie_colors)
            xlsx = _build_xlsx(answers, [lbl], [values])
            _embed_chart(doc, chart_xml, xlsx, _next_part(), cy=PIE_H)
        _add_caption(
            f'Рисунок {n} – Распределение ответов респондентов на вопрос: «{question_name}»',
            above=False
        )


def _insert_word_table(doc, q: dict):
    """Insert a full-width Word table for this question's data."""
    file_keys   = q['file_keys']
    file_labels = q['file_labels']
    rows_data   = q['rows']
    file_totals = q['file_totals']
    show_total  = q.get('show_total', True)
    hidden_col  = q.get('hidden_col', 'none')
    is_single   = len(file_keys) == 1

    # column count
    show_count = hidden_col != 'count'
    show_pct   = hidden_col != 'percent'
    if is_single:
        n_cols = 1 + show_count + show_pct
    else:
        n_cols = 1 + len(file_keys) * show_count + len(file_keys) * show_pct

    n_rows = 1 + len(rows_data) + (1 if show_total else 0)
    table = doc.add_table(rows=n_rows, cols=n_cols)

    # full-width + borders via XML
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    borders_el = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        borders_el.append(b)
    tblPr.append(borders_el)

    def cell_text(cell, text, bold=False, center=False):
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(str(text))
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── header ──────────────────────────────────────────────────────────────
    hdr = table.rows[0].cells
    ci = 0
    cell_text(hdr[ci], q.get('h1', 'Ответ'), bold=True); ci += 1

    if is_single:
        if show_count:
            cell_text(hdr[ci], q.get('h2', 'Кол-во ответивших'), bold=True, center=True); ci += 1
        if show_pct:
            cell_text(hdr[ci], q.get('h3', '% от числа ответивших'), bold=True, center=True); ci += 1
    else:
        if show_count:
            for fk in file_keys:
                cell_text(hdr[ci], f'{file_labels.get(fk, fk)}, кол-во', bold=True, center=True); ci += 1
        if show_pct:
            for fk in file_keys:
                cell_text(hdr[ci], f'{file_labels.get(fk, fk)}, %', bold=True, center=True); ci += 1

    # ── data rows ───────────────────────────────────────────────────────────
    for ri, row in enumerate(rows_data):
        cells = table.rows[ri + 1].cells
        ci = 0
        cell_text(cells[ci], row['answer']); ci += 1

        if is_single:
            fk = file_keys[0]
            total = file_totals.get(fk, 0)
            count = row['counts'].get(fk, 0)
            if show_count:
                cell_text(cells[ci], count, center=True); ci += 1
            if show_pct:
                pct = f'{count / total * 100:.1f}%' if total else '—'
                cell_text(cells[ci], pct, center=True); ci += 1
        else:
            if show_count:
                for fk in file_keys:
                    cell_text(cells[ci], row['counts'].get(fk, 0), center=True); ci += 1
            if show_pct:
                for fk in file_keys:
                    count = row['counts'].get(fk, 0)
                    total = file_totals.get(fk, 0)
                    pct = f'{count / total * 100:.1f}%' if total else '—'
                    cell_text(cells[ci], pct, center=True); ci += 1

    # ── total row ───────────────────────────────────────────────────────────
    if show_total:
        cells = table.rows[-1].cells
        ci = 0
        cell_text(cells[ci], 'Всего', bold=True); ci += 1

        if is_single:
            fk = file_keys[0]
            if show_count:
                cell_text(cells[ci], file_totals.get(fk, 0), bold=True, center=True); ci += 1
            if show_pct:
                cell_text(cells[ci], '100%', bold=True, center=True); ci += 1
        else:
            if show_count:
                for fk in file_keys:
                    cell_text(cells[ci], file_totals.get(fk, 0), bold=True, center=True); ci += 1
            if show_pct:
                for fk in file_keys:
                    cell_text(cells[ci], '100%', bold=True, center=True); ci += 1
